import os
import pandas as pd
from document_processor import DocumentProcessorFactory

# Create a directory for test files if it doesn't exist
test_dir = "test_files"
os.makedirs(test_dir, exist_ok=True)

# Sample CSV content
csv_content = """Name,Age,Occupation
Alice,30,Engineer
Bob,25,Designer
Charlie,35,Teacher
"""

# Sample Excel content
excel_data = {
    "Name": ["Alice", "Bob", "Charlie"],
    "Age": [30, 25, 35],
    "Occupation": ["Engineer", "Designer", "Teacher"]
}
excel_file_path = os.path.join(test_dir, "test.xlsx")
pd.DataFrame(excel_data).to_excel(excel_file_path, index=False)

# Sample PDF content (you can create a simple PDF with a table using a library like ReportLab)
# For simplicity, we will assume you have a PDF file named 'test.pdf' in the test_files directory.

# Sample DOCX content
from docx import Document
docx_file_path = os.path.join(test_dir, "test.docx")
doc = Document()
doc.add_heading('Test Document', level=1)
doc.add_paragraph('This is a test document for the Document Processor.')
doc.add_paragraph('Name, Age, Occupation')
doc.add_paragraph('Alice, 30, Engineer')
doc.add_paragraph('Bob, 25, Designer')
doc.add_paragraph('Charlie, 35, Teacher')
doc.save(docx_file_path)

# Sample image file (you can create a simple image with a table using OpenCV or PIL)
# For simplicity, we will assume you have an image file named 'test_image.png' in the test_files directory.

# Sample text file
text_file_path = os.path.join(test_dir, "test.txt")
with open(text_file_path, 'w') as f:
    f.write(csv_content)

# Function to test document processing
def test_document_processing(file_path):
    processor = DocumentProcessorFactory.get_processor(file_path)
    df = processor.process(file_path)
    print(f"Processed {file_path}:")
    print(df)

# Test each file type
test_files = [
    os.path.join(test_dir, "test.txt"),
    os.path.join(test_dir, "test.xlsx"),
    os.path.join(test_dir, "test.docx"),
    # Add paths to your PDF and image files here
]

for test_file in test_files:
    test_document_processing(test_file)
